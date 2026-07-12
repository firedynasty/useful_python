#!/usr/bin/env python3
"""
Enhanced HTML to DOCX converter that preserves:
- Tables with borders and styling
- Text formatting (bold, italic, underline)
- Lists (ordered and unordered)
- Headings (h1-h6)
- Links
- Text colors and highlighting
- Alignment
"""

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import argparse


def parse_css_color(color_str):
    """Parse CSS color string to RGB tuple."""
    if not color_str:
        return None
    
    # Handle hex colors
    if color_str.startswith('#'):
        color_str = color_str.lstrip('#')
        if len(color_str) == 3:
            color_str = ''.join([c*2 for c in color_str])
        if len(color_str) == 6:
            return RGBColor(
                int(color_str[0:2], 16),
                int(color_str[2:4], 16),
                int(color_str[4:6], 16)
            )
    
    # Handle rgb() format
    rgb_match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_str)
    if rgb_match:
        return RGBColor(
            int(rgb_match.group(1)),
            int(rgb_match.group(2)),
            int(rgb_match.group(3))
        )
    
    # Handle named colors (basic set)
    named_colors = {
        'red': RGBColor(255, 0, 0),
        'blue': RGBColor(0, 0, 255),
        'green': RGBColor(0, 128, 0),
        'black': RGBColor(0, 0, 0),
        'white': RGBColor(255, 255, 255),
        'gray': RGBColor(128, 128, 128),
        'grey': RGBColor(128, 128, 128),
    }
    
    return named_colors.get(color_str.lower())


def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Check for each border
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '#000000'))
            tcPr.append(element)


def parse_style_attribute(style_str):
    """Parse HTML style attribute into a dictionary."""
    styles = {}
    if not style_str:
        return styles
    
    for item in style_str.split(';'):
        if ':' in item:
            key, value = item.split(':', 1)
            styles[key.strip().lower()] = value.strip()
    
    return styles


def process_inline_elements(element, paragraph, run=None):
    """Process inline HTML elements and add them to a paragraph."""
    if run is None:
        run = paragraph.add_run()
    
    # Handle text nodes
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip():
            run.text = text
            return run
        return None
    
    # Get style attributes
    styles = parse_style_attribute(element.get('style', ''))
    
    # Process based on tag
    tag = element.name.lower()
    
    if tag in ['b', 'strong']:
        run.bold = True
    elif tag in ['i', 'em']:
        run.italic = True
    elif tag in ['u', 'ins']:
        run.underline = True
    elif tag == 'span':
        # Apply any inline styles
        if 'color' in styles:
            color = parse_css_color(styles['color'])
            if color:
                run.font.color.rgb = color
        if 'background-color' in styles or 'background' in styles:
            # Note: python-docx doesn't support background colors directly
            # You'd need to use highlighting instead
            pass
        if 'font-size' in styles:
            size_str = styles['font-size']
            if 'pt' in size_str:
                size = float(size_str.replace('pt', ''))
                run.font.size = Pt(size)
        if 'font-weight' in styles and styles['font-weight'] in ['bold', '700']:
            run.bold = True
        if 'font-style' in styles and styles['font-style'] == 'italic':
            run.italic = True
    elif tag == 'a':
        # For links, just add the text in blue and underlined
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.underline = True
    elif tag == 'code':
        run.font.name = 'Courier New'
    elif tag == 'br':
        run.add_break()
        return run
    
    # Process children
    for child in element.children:
        if isinstance(child, NavigableString):
            if str(child).strip():
                if not run.text:
                    run.text = str(child)
                else:
                    run.text += str(child)
        else:
            # Recursively process inline elements
            process_inline_elements(child, paragraph, run)
    
    return run


def add_html_table_to_doc(table_element, doc):
    """Convert HTML table to Word table with styling."""
    # Count rows and columns
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Find maximum number of columns
    max_cols = 0
    for row in rows:
        cells = row.find_all(['td', 'th'])
        col_count = sum(int(cell.get('colspan', 1)) for cell in cells)
        max_cols = max(max_cols, col_count)
    
    if max_cols == 0:
        return
    
    # Create table
    word_table = doc.add_table(rows=0, cols=max_cols)
    
    # Apply table styling
    table_style = table_element.get('style', '')
    styles = parse_style_attribute(table_style)
    
    # Check for borders
    has_border = 'border' in styles or table_element.get('border', '0') != '0'
    
    if has_border:
        word_table.style = 'Table Grid'
    
    # Process each row
    for row_element in rows:
        word_row = word_table.add_row()
        cells = row_element.find_all(['td', 'th'])
        
        col_index = 0
        for cell_element in cells:
            if col_index >= max_cols:
                break
            
            # Get cell content
            word_cell = word_row.cells[col_index]
            
            # Clear default paragraph
            word_cell.text = ''
            
            # Process cell content
            for child in cell_element.children:
                if child.name in ['p', 'div'] or isinstance(child, NavigableString):
                    p = word_cell.add_paragraph()
                    
                    # Handle alignment
                    cell_style = parse_style_attribute(cell_element.get('style', ''))
                    align = cell_element.get('align') or cell_style.get('text-align')
                    if align:
                        if align == 'center':
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif align == 'right':
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif align == 'justify':
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Process content
                    if isinstance(child, NavigableString):
                        if str(child).strip():
                            run = p.add_run(str(child))
                            # Bold for header cells
                            if cell_element.name == 'th':
                                run.bold = True
                    else:
                        for elem in child.children:
                            process_inline_elements(elem, p)
                elif child.name in ['ul', 'ol']:
                    # Handle lists inside table cells
                    process_list(child, word_cell)
            
            # Apply cell borders if needed
            if has_border:
                border_props = {"sz": 8, "val": "single", "color": "#000000"}
                set_cell_border(word_cell, top=border_props, left=border_props, 
                              bottom=border_props, right=border_props)
            
            # Handle colspan
            colspan = int(cell_element.get('colspan', 1))
            if colspan > 1:
                for i in range(1, colspan):
                    if col_index + i < max_cols:
                        # Merge cells
                        word_cell.merge(word_row.cells[col_index + i])
            
            col_index += colspan
    
    # Add spacing after table
    doc.add_paragraph()


def process_list(list_element, parent):
    """Process HTML list (ul/ol) and add to document."""
    is_ordered = list_element.name == 'ol'
    items = list_element.find_all('li', recursive=False)
    
    for i, item in enumerate(items, 1):
        # Create paragraph for list item
        if hasattr(parent, 'add_paragraph'):
            p = parent.add_paragraph()
        else:
            # For table cells
            p = parent.paragraphs[-1] if parent.paragraphs else parent.add_paragraph()
            p = parent.add_paragraph()
        
        # Add bullet or number
        if is_ordered:
            prefix = f"{i}. "
        else:
            prefix = "• "
        
        run = p.add_run(prefix)
        
        # Process item content
        for child in item.children:
            if isinstance(child, NavigableString):
                if str(child).strip():
                    p.add_run(str(child))
            elif child.name in ['ul', 'ol']:
                # Nested list
                process_list(child, parent)
            else:
                process_inline_elements(child, p)


def html_to_docx(html_file, output_file=None, encoding='utf-8'):
    """
    Convert HTML file to DOCX format.
    
    Args:
        html_file: Path to input HTML file
        output_file: Path to output DOCX file (optional, defaults to input name with .docx)
        encoding: HTML file encoding (default: utf-8)
    
    Returns:
        Path to created DOCX file
    """
    # Read HTML
    with open(html_file, 'r', encoding=encoding) as f:
        html_content = f.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create document
    doc = Document()
    
    # Extract title from HTML if present
    title = soup.find('title')
    if title:
        doc.core_properties.title = title.text.strip()
    
    # Process body content (or entire soup if no body tag)
    body = soup.find('body') or soup
    
    # Process all elements
    for element in body.children:
        if isinstance(element, NavigableString):
            # Skip pure whitespace
            if str(element).strip():
                doc.add_paragraph(str(element))
            continue
        
        tag = element.name.lower()
        
        # Handle different element types
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag[1])
            heading = doc.add_heading(level=min(level, 9))
            heading.text = element.get_text().strip()
            
        elif tag == 'p':
            p = doc.add_paragraph()
            # Parse style for alignment
            styles = parse_style_attribute(element.get('style', ''))
            align = element.get('align') or styles.get('text-align')
            if align:
                if align == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == 'justify':
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Process paragraph content
            for child in element.children:
                process_inline_elements(child, p)
        
        elif tag == 'div':
            # Treat div as paragraph
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'table':
                    add_html_table_to_doc(child, doc)
                elif child.name in ['ul', 'ol']:
                    process_list(child, doc)
                else:
                    process_inline_elements(child, p)
        
        elif tag == 'table':
            add_html_table_to_doc(element, doc)
        
        elif tag in ['ul', 'ol']:
            process_list(element, doc)
        
        elif tag == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            for child in element.children:
                process_inline_elements(child, p)
        
        elif tag == 'pre':
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.font.name = 'Courier New'
        
        elif tag == 'hr':
            # Add a horizontal line (paragraph with border)
            p = doc.add_paragraph()
            p.add_run('_' * 50)
    
    # Save document
    if output_file is None:
        output_file = html_file.rsplit('.', 1)[0] + '.docx'
    
    doc.save(output_file)
    print(f"Successfully converted '{html_file}' to '{output_file}'")
    return output_file


def main():
    parser = argparse.ArgumentParser(description='Convert HTML file to DOCX with formatting preservation')
    parser.add_argument('input', help='Input HTML file path')
    parser.add_argument('-o', '--output', help='Output DOCX file path (default: input_name.docx)')
    parser.add_argument('-e', '--encoding', default='utf-8', help='HTML file encoding (default: utf-8)')
    
    args = parser.parse_args()
    
    try:
        html_to_docx(args.input, args.output, args.encoding)
    except Exception as e:
        print(f"Error: {e}")
        return 1
    
    return 0


if __name__ == '__main__':
    exit(main())
