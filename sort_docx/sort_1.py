"""
Sort a Word document (.docx) by headings while preserving formatting.

Usage:
    python sort_docx.py input.docx output.docx [heading_level]
    
Examples:
    python sort_docx.py input.docx output.docx 3    # Sort by Heading 3
    python sort_docx.py input.docx output.docx      # Sort by Heading 1 (default)

Requirements:
    pip install python-docx
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
import sys
import re

def get_heading_level(paragraph):
    """Get the heading level of a paragraph (1-9), or 0 if not a heading."""
    if paragraph.style.name.startswith('Heading'):
        match = re.search(r'\d+', paragraph.style.name)
        if match:
            return int(match.group())
    return 0

def get_heading_text(paragraph):
    """Extract plain text from a heading paragraph."""
    return paragraph.text.strip()

def copy_paragraph(source_para, target_doc):
    """Copy a paragraph with all its formatting and hyperlinks to the target document."""
    # Create a new paragraph
    new_para = target_doc.add_paragraph()
    
    # Deep copy the entire paragraph XML element to preserve everything
    new_para._element.getparent().replace(new_para._element, deepcopy(source_para._element))
    
    return new_para

def copy_table(source_table, target_doc):
    """Copy a table to the target document."""
    # Deep copy the table XML element
    new_table_element = deepcopy(source_table._element)
    target_doc.element.body.append(new_table_element)
    
    # Return the table object (find it in the document's tables list)
    return target_doc.tables[-1]

def extract_sections_by_level(doc, target_level):
    """Extract sections grouped by specified heading level."""
    sections = []
    current_section = None
    
    for element in doc.element.body:
        if element.tag.endswith('p'):
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if para:
                level = get_heading_level(para)
                
                if level == target_level:
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        'heading': para,
                        'heading_text': get_heading_text(para),
                        'content': []
                    }
                elif current_section is not None:
                    current_section['content'].append(('paragraph', para))
                else:
                    if not sections or sections[-1].get('heading') is not None:
                        sections.append({
                            'heading': None,
                            'heading_text': '',
                            'content': []
                        })
                    sections[-1]['content'].append(('paragraph', para))
        
        elif element.tag.endswith('tbl'):
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            
            if table:
                if current_section is not None:
                    current_section['content'].append(('table', table))
                else:
                    if not sections or sections[-1].get('heading') is not None:
                        sections.append({
                            'heading': None,
                            'heading_text': '',
                            'content': []
                        })
                    sections[-1]['content'].append(('table', table))
    
    if current_section:
        sections.append(current_section)
    
    return sections

def sort_document(input_path, output_path, heading_level=1):
    """Sort a Word document by specified heading level sections alphabetically."""
    doc = Document(input_path)
    sections = extract_sections_by_level(doc, heading_level)
    
    sections_with_headings = [s for s in sections if s['heading'] is not None]
    sections_without_headings = [s for s in sections if s['heading'] is None]
    
    sections_with_headings.sort(key=lambda x: x['heading_text'].lower())
    
    new_doc = Document()
    
    for section in sections_without_headings:
        for content_type, content in section['content']:
            if content_type == 'paragraph':
                copy_paragraph(content, new_doc)
            elif content_type == 'table':
                copy_table(content, new_doc)
    
    for section in sections_with_headings:
        copy_paragraph(section['heading'], new_doc)
        
        for content_type, content in section['content']:
            if content_type == 'paragraph':
                copy_paragraph(content, new_doc)
            elif content_type == 'table':
                copy_table(content, new_doc)
    
    new_doc.save(output_path)
    print(f"Sorted document saved to: {output_path}")
    print(f"Sorted {len(sections_with_headings)} sections by Heading {heading_level}")

if __name__ == "__main__":
    if len(sys.argv) < 3 or len(sys.argv) > 4:
        print("Usage: python sort_docx.py input.docx output.docx [heading_level]")
        print("Example: python sort_docx.py input.docx output.docx 3")
        print("Default heading level is 1 if not specified")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    heading_level = int(sys.argv[3]) if len(sys.argv) == 4 else 1
    
    if heading_level < 1 or heading_level > 9:
        print("Error: Heading level must be between 1 and 9")
        sys.exit(1)
    
    sort_document(input_file, output_file, heading_level)
